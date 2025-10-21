import docx
import json
import re

from datetime import datetime

from src.generate_table import generate_table
from src.add_hyperlink import add_hyperlink

#dbt_path = "/home/daniele/dbt/credem/hubble/xanage-edp"
model_types = {
    "Hub-Chiavi": "Per tutte le tabelle hub sono presenti i test di univocità sia della chiave naturale sia della chiave di hash.",
    "Link-Relazioni": "Per tutte le tabelle link sono presenti i test di univocità sia della chiave naturale sia della chiave di hash.",
    "Reference-Domini/Non_storici": "Per tutte le tabelle reference non storiche sono presenti i test di univocità della chiave naturale.",
    "Reference-Domini/Storici_delta": None,
    "Reference-Domini/Storici_full": None,
    "Satellite-Attributi/Storici_delta": None,
    "Satellite-Attributi/Storici_full": None,
}


def generate_doc(dbt_path: str):

    with open(f"{dbt_path}/target/manifest.json", "r") as f:
        manifest = json.load(f)
    prj_name = manifest["metadata"]["project_name"]
    nodes = manifest["nodes"]
    doc_title = f"{prj_name.upper()} HUBBLE dbt"

    doc = docx.Document("templates/template.docx")

    doc_properties = doc.core_properties
    doc_properties.title = doc_title
    doc_properties.author = "sdg"

    doc.add_heading("Informazioni Documento", level=1)
    doc.add_heading("Cronologia", level=2)
    generate_table(
      doc=doc,
      header=["Versione", "Data", "Descrizione", "Autore"],
      body=[
        ["1.0", datetime.today().strftime("%d/%m/%Y"), "Analisi tecnica", "sdg"],
      ]
    )
                  
    doc.add_heading("Validazione documento", level=2)
    generate_table(
      doc=doc,
      header=["Responsabile Valutazione", "Unità Organizzativa", "Data Approvazione", "Stato"],
      body=[
        ["", "", "", ""],
      ]
    )

    doc.add_page_break()

    doc.add_heading("Glossario", level=1)
    doc.add_paragraph(("Inserire un riferimento ai termini e sigle utilizzati al fine "
                      "di eliminare/minimizzare condizioni di ambiguità terminologiche "
                      "inerenti sia aspetti tecnici (es. Token OTP), sia aspetti correlabili "
                      "a prodotti/oggetti tipici del mondo bancario (es. Catalogo Prodotti, "
                      "Conto Deposito etc..) e ad aspetti di natura organizzativa "
                      "(es. ruoli specifici presenti nei processi analizzati).\n"))
    generate_table(
      doc=doc,
      header=["ID", "Termine", "Descrizione"],
      body=[
        ["", "", ""],
      ]
    )

    doc.add_page_break()

    doc.add_heading("Premessa", level=1)
    doc.add_heading("Flussi", level=2)
    doc.add_paragraph(("E’ stato richiesto di replicare in ambiente Cloud "
                      "dbt/BigQuery i caricamenti dei seguenti flussi:\n"))
    # TODO read from pcs_anag_flussi seed file
    generate_table(
      doc=doc,
      header=["Flusso", "Periodicità", "Tipo flusso", "Tipo input", "Note"],
      body=[
        ["", "", "", "", ""],
      ]
    )
    doc.add_paragraph(f"\ne in particolare nella Subject Area {prj_name.upper()}.\n")
    doc.add_paragraph(("Lo sviluppo ha come obiettivo il porting su dbt delle "
                        "attuali logiche di caricamento presenti su DataStage e "
                        "sulle query del modello Data Vault salvate nelle tabelle "
                        "Teradata del database PCS.\n"))
    doc.add_paragraph(("Nei capitoli seguenti vengono dettagliati i componenti "
                      "sviluppati nell’ambito del progetto."))

    doc.add_page_break()

    doc.add_heading("Back-End - HUBBLE", level=1)
    paragraph = doc.add_paragraph()
    paragraph.text = "Per i dettagli relativi alle diverse componenti si rimanda al "
    add_hyperlink(paragraph=paragraph, text="documento di progettazione.", url="https://docs.google.com/document/u/1/d/1F9nKcThgW3IIWqGHxDgSoa3FUj_K23f1FzQHJ0ZLhrU/")
    doc.add_paragraph(("I modelli dbt, le definizioni delle external table e "
                      "dei test di seguito menzionati sono consultabili all’interno "
                      "del repository Azure DevOps xxxx."))
    doc.add_paragraph("Il progetto utilizza inoltre i seguenti package dbt:")
    # TODO read dynamically the list pf packages from manifest
    doc.add_paragraph("macro_dbt_hublle (repository Azure DevOps DBTELT/dbtelt-dbtmacro)\n", style='List Paragraph1')
    generate_table(
      doc=doc,
      header=["ID Componente", "Nome Componente"],
      body=[
        ["1", "bronze - External Table + Staging"],
        ["2", "silver - Data Vault (HUB, Satelliti, Link, Reference)"],
        ["3", "gold - Outputs"],
      ]
    )
    doc.add_paragraph(("\nL’intero processo si articola attraverso i seguenti dataset "
                      "del progetto GCP prj-gcp-crg-xanage-edp-01:"))
    doc.add_paragraph("bronze: contenente le external table e le viste di staging HS", style='List Paragraph1')
    doc.add_paragraph("silver: contenente il modello DataVault (tabelle HUB, LINK, SATELLITI, REFERENCE)", style='List Paragraph1')
    doc.add_paragraph("gold: contenente le viste / tabelle di output", style='List Paragraph1')
    doc.add_paragraph()

    doc.add_heading("Bronze Layer", level=2)
    doc.add_paragraph(("Le external table presentano il contenuto dei file EBCDIC "
                      "(convertiti in formato parquet tramite apposito cloud job) "
                      "senza alcuna trasformazione; ogni file ha una corrispondente "
                      "external table, denominata come il file, con l’eccezione dei "
                      "file header che sono raggruppati sotto un’unica external "
                      "table denominata HEADER.\n"))
    doc.add_paragraph(("Le viste di staging sono anch’esse una per ogni file EBCDIC, "
                      "denominate HS_<nome file>, e contengono le trasformazioni di "
                      "data type ed i controlli di formato comuni a tutti i successivi "
                      "modelli Data Vault, oltre alla rinomina di tutti i campi "
                      "secondo lo standard del modello dati EDP ed alla creazione "
                      "delle chiavi di hash (ove necessario).\n"))

    doc.add_heading("Trasformazioni applicate", level=3)
    doc.add_paragraph("Di seguito le trasformazioni applicate:")
    doc.add_paragraph("check_date_format: validazione dei campi data", 
                      style='List Paragraph1')
    doc.add_paragraph("check_timestamp_format: validazione dei campi timestamp", 
                      style='List Paragraph1')
    doc.add_paragraph("check_hash_format: validazione del formato dei campi da concatenare nella costruzione di hash keys", 
                      style='List Paragraph1')
    doc.add_paragraph("rimozione degli zeri non significativi nei codici CDG ed istituto.", 
                      style='List Paragraph1')
    doc.add_paragraph(("\nPer i dettagli implementativi si veda la documentazione "
                      "del package “macro_dbt_hubble”.\n"))

    doc.add_heading("Regole di Data Quality", level=3)
    doc.add_paragraph("Di seguito le regole applicate:")
    doc.add_paragraph(("controllo di chiave primaria su tutte le viste di staging, "
                      "applicato sui LOAD_TS degli ultimi 3 giorni di calendario"),
                        style='List Paragraph1')
    doc.add_paragraph()

    doc.add_heading("Hash Key", level=3)
    hash_keys = []
    for n in nodes.values():
        if n["resource_type"] == "model" and "Staging" in n["path"] and "hash_key" in n["raw_code"]:
            hks = re.findall(r"field_prefix='(\w+)", n["raw_code"])
            for hk in hks:
                hash_keys.append([n["name"], f"{hk}_HK"])
    generate_table(
      doc=doc,
      header=["Tabella", "Hash Key"],
      body=hash_keys,
    )

    wrk_tables = []
    for n in nodes.values():
        if n["resource_type"] == "model" and "Work" in n["path"]:
            wrk_tables.append([n["name"], "\n".join(n["depends_on"]["nodes"]).replace(f"model.{prj_name}.HS_", "")])
    if len(wrk_tables) > 0:
        doc.add_heading("Tabelle di WORK", level=3)
        paragraph = doc.add_paragraph()
        font = paragraph.add_run("... qui spiegare perchè siano necessarie queste tabelle di work ...").font
        #font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        paragraph = doc.add_paragraph()
        generate_table(
          doc=doc,
          header=["Tabella target", "Flusso"],
          body=wrk_tables,
        )

    doc.add_heading("Silver Layer", level=2)
    doc.add_paragraph("...")
    for model_type in model_types.keys():
        data_vault_tables = []
        for n in nodes.values():
            if n["resource_type"] == "model" and model_type in n["path"]:
                data_vault_tables.append([
                    n["name"],
                  "\n".join(n["depends_on"]["nodes"]). \
                        replace(f"model.{prj_name}.HS_", ""). \
                        replace(f"model.{prj_name}.WRK_{n["name"].split("_")[0]}_", "")
                  ])
        if len(data_vault_tables) > 0:
            doc.add_heading(f"Caricamento {model_type.replace("/", " - ").replace("_", " ")}", level=3)
            doc.add_paragraph(("Di seguito l’elenco degli step di caricamento delle "
                              f"tabelle {model_type.split("-")[0]}, i cui relativi "
                              "modelli possono essere consultati nella cartella "
                              f"models/{model_type} del progetto dbt:"))
            generate_table(
              doc=doc,
              header=["Tabella target", "Flusso"],
              body=data_vault_tables,
            )
            if model_types[model_type] is not None:
                doc.add_heading("Regole di Data Quality", level=4)
                doc.add_paragraph(model_types[model_type])

    doc.add_heading("Gold Layer", level=2)
    doc.add_paragraph("Di seguito l’elenco delle viste implementate nel gold layer:")
    gold_tables = []
    for n in nodes.values():
        if n["resource_type"] == "model" and "gold" in n["path"].lower():
            gold_tables.append([n["name"], "\n".join(n["depends_on"]["nodes"]).replace(f"model.{prj_name}.", "")])
    generate_table(
      doc=doc,
      header=["Tabella target", "Sorgenti"],
      body=gold_tables,
    )

    doc.save(f"{prj_name.upper()}_dbt_ATE.docx")
